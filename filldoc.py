from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH


SCORE_RANGE = "الكفاءة"
CATEGORY_RATION = "نسبة الفئة"

class FillDoc:
    def __init__(self, output_path, data, datatable, datapicture):
        self.output_path = output_path
        self.data = data
        self.datatable = datatable
        self.datapicture = datapicture
        self.doc = Document("./Doc2.docx")
        self.filldocstr()

    def filldocstr(self):
        for paragraph in self.doc.paragraphs:
            for key, value in self.data.items():
                if key in paragraph.text:
                    for run in paragraph.runs:
                        run.text = run.text.replace(key, str(value))
        # for paragraph in doc.paragraphs:
            if "[percentage_distribution]" in paragraph.text:
                paragraph.text = paragraph.text.replace("[percentage_distribution]", '')
                # Crate Table
                t = self.doc.add_table(rows=2, cols=5)
                t.autofit = False

                # Row Category
                row_category = t.rows[0].cells
                row_category[4].text = SCORE_RANGE
                # Row percentage
                row_percentage = t.rows[1].cells
                row_percentage[4].text = CATEGORY_RATION

                i = 0
                for index, value in self.datatable.items():
                    col = t.columns[i].cells
                    col[0].text = index
                    col[1].text = str(round(value, 2))
                    i += 1
                t_element = t._tbl
                p_element = paragraph._element
                p_element.addnext(t_element)
                paragraph.clear()

            for key, value in self.datapicture.items():
                if key in paragraph.text:
                    for run in paragraph.runs:
                        paragraph.text = paragraph.text.replace(key, "")
                        run = paragraph.add_run()
                        run.add_picture(value+".png", width=Cm(10))
                        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

        self.doc.save(self.output_path)







